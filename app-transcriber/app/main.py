import asyncio
import io
import os
import re

import yt_dlp
from docx import Document
from fastapi import Depends, FastAPI, File, Form, Request, UploadFile
from fastapi.responses import HTMLResponse, JSONResponse, Response, StreamingResponse
from fastapi.templating import Jinja2Templates

from .auth import verify_token
from .dao.transcription_dao import TranscriptionDAO
from .services import subtitles as subs_service
from .services import whisper_service

app = FastAPI(root_path="/transcriber")
_APP_DIR = os.path.dirname(os.path.abspath(__file__))
templates = Jinja2Templates(directory=os.path.join(_APP_DIR, "templates"))
templates.env.globals["getattr"] = getattr
templates.env.globals["portal_url"] = "/"
templates.env.globals["favicon_url"] = "/static/favicon.png"
templates.env.globals["dashboard_url"] = "/"

COOKIES_FILE = "/app/data/cookies.txt"
os.makedirs("/app/data", exist_ok=True)


def _format_duration(seconds):
    """Formato de duración: 65 -> '1:05', 3665 -> '1:01:05', None -> '—'."""
    if seconds is None or (isinstance(seconds, (int, float)) and seconds <= 0):
        return "—"
    try:
        s = int(round(float(seconds)))
    except (TypeError, ValueError):
        return "—"
    if s < 60:
        return f"0:{s:02d}"
    m, s = divmod(s, 60)
    if m < 60:
        return f"{m}:{s:02d}"
    h, m = divmod(m, 60)
    return f"{h}:{m:02d}:{s:02d}"


templates.env.filters["format_duration"] = _format_duration

YDL_OPTS_BASE = {
    "quiet": True,
    "no_warnings": True,
    "user_agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
}


def _get_ydl_opts(base_opts: dict = None) -> dict:
    """Retorna opciones de yt-dlp con cookies si existen."""
    opts = {**YDL_OPTS_BASE, **(base_opts or {})}
    if os.path.isfile(COOKIES_FILE):
        opts["cookiefile"] = COOKIES_FILE
    return opts


def _is_youtube_url(url: str) -> bool:
    return bool(url and ("youtube.com" in url or "youtu.be" in url))


SUBTITLES_TIMEOUT = 45
WHISPER_TIMEOUT = 480  # 8 min para transcripción síncrona
MAX_DURATION_WHISPER = 600  # 10 min máximo para Whisper sin subtítulos (más → job en background)

# Mismo valor que en whisper_service para dividir audio largo
CHUNK_DURATION_SEC = 600  # 10 min por segmento


def _sync_transcription_worker(
    job_id: int,
    url: str,
    lang: str,
    include_ts: bool,
    model: str,
):
    """Blocking pipeline: try subtitles, else download + Whisper (por trozos si > 10 min). Updates DAO."""
    dao = TranscriptionDAO()
    try:
        transcript, _ = subs_service.get_subtitles(
            url, language=lang, include_timestamps=include_ts
        )
        if transcript and (transcript or "").strip():
            dao.update_completed(job_id, transcript, "subtitles", duration_seconds=None)
            return
    except Exception:
        pass
    audio_path = None
    files_to_remove = []
    try:
        audio_path, err_audio = whisper_service.download_audio(url, COOKIES_FILE)
        if not audio_path:
            dao.update_failed(job_id, err_audio or "No se pudo descargar el audio.")
            return
        files_to_remove.append(audio_path)

        duration = whisper_service.get_audio_duration(audio_path)
        chunks = whisper_service.split_audio_into_chunks(
            audio_path, chunk_sec=CHUNK_DURATION_SEC
        )

        if not chunks:
            dao.update_failed(job_id, "No se pudo procesar el audio.")
            return

        if len(chunks) == 1 and chunks[0] == audio_path:
            # Audio corto: transcribir todo de una
            transcript = whisper_service.transcribe_audio(
                audio_path,
                model_size=model,
                language=lang if lang != "auto" else None,
                include_timestamps=include_ts,
            )
        else:
            # Audio largo: cada elemento de chunks es un archivo distinto (chunk_0000.mp3, ...)
            for p in chunks:
                if p != audio_path:
                    files_to_remove.append(p)
            parts = []
            for i, chunk_path in enumerate(chunks):
                part = whisper_service.transcribe_audio(
                    chunk_path,
                    model_size=model,
                    language=lang if lang != "auto" else None,
                    include_timestamps=include_ts,
                    time_offset_sec=i * CHUNK_DURATION_SEC,
                )
                if (part or "").strip():
                    parts.append(part)
            transcript = "\n\n".join(parts).strip() if parts else ""

        if (transcript or "").strip():
            dao.update_completed(job_id, transcript, "whisper", duration_seconds=duration)
        else:
            dao.update_failed(job_id, "La transcripción con Whisper no generó texto.")
    except Exception as e:
        dao.update_failed(job_id, str(e))
    finally:
        for p in files_to_remove:
            try:
                if p and os.path.isfile(p):
                    os.remove(p)
            except OSError:
                pass


async def _run_transcription_job(
    job_id: int,
    url: str,
    lang: str,
    include_ts: bool,
    model: str,
):
    """Run blocking worker in thread pool."""
    await asyncio.to_thread(
        _sync_transcription_worker,
        job_id, url, lang, include_ts, model,
    )


@app.get("/", response_class=HTMLResponse)
async def index(request: Request, user=Depends(verify_token)):
    dao = TranscriptionDAO()
    history = dao.get_recent(limit=50)
    dashboard_url = "/dashboard"
    return templates.TemplateResponse(
        "index.html",
        {"request": request, "history": history, "user": user, "dashboard_url": dashboard_url},
    )


@app.post("/preview")
async def preview(
    url: str = Form(...),
    user=Depends(verify_token),
):
    url = (url or "").strip()
    if not url:
        return JSONResponse({"error": "Escribe una URL de YouTube."}, status_code=400)
    if not _is_youtube_url(url):
        return JSONResponse(
            {"error": "Solo se admiten enlaces de YouTube (youtube.com o youtu.be)."},
            status_code=400,
        )
    opts = _get_ydl_opts({"format": "best"})
    try:
        with yt_dlp.YoutubeDL(opts) as ydl:
            info = await asyncio.to_thread(ydl.extract_info, url, download=False)
        if not info:
            return JSONResponse(
                {"error": "No se pudo obtener el vídeo. Comprueba que el enlace sea correcto y público."},
                status_code=400,
            )
        subtitles = info.get("subtitles") or info.get("automatic_captions") or {}
        sub_langs = list(subtitles.keys()) if isinstance(subtitles, dict) else []
        return JSONResponse({
            "title": info.get("title"),
            "duration": info.get("duration"),
            "uploader": info.get("uploader") or info.get("channel"),
            "thumbnail": info.get("thumbnail"),
            "subtitle_languages": sub_langs[:20],
        })
    except Exception as e:
        return JSONResponse(
            {"error": (str(e).split("\n")[0] if str(e) else "Error al obtener el vídeo.")},
            status_code=400,
        )


@app.post("/transcribir")
async def transcribir(
    request: Request,
    url: str = Form(...),
    language: str = Form("auto"),
    include_timestamps: str = Form("true"),
    whisper_model: str = Form("small"),
    user=Depends(verify_token),
):
    url = (url or "").strip()
    if not url or not _is_youtube_url(url):
        return JSONResponse({"error": "URL de YouTube no válida."}, status_code=400)
    include_ts = include_timestamps.lower() in ("true", "1", "yes", "on")
    model = whisper_model if whisper_model in ("base", "small", "medium", "large-v2", "large-v3") else "small"
    lang = (language or "auto").strip() or "auto"

    # 0) Obtener info del vídeo (título y duración) — siempre útil
    video_info = {}
    try:
        opts = _get_ydl_opts()
        with yt_dlp.YoutubeDL(opts) as ydl:
            video_info = await asyncio.wait_for(
                asyncio.to_thread(ydl.extract_info, url, download=False),
                timeout=30,
            ) or {}
    except Exception:
        pass
    video_title = video_info.get("title") or ""
    video_duration = video_info.get("duration") or 0  # segundos
    subtitles_raw = video_info.get("subtitles") or video_info.get("automatic_captions") or {}
    has_subtitles = bool(subtitles_raw and isinstance(subtitles_raw, dict) and list(subtitles_raw))

    # Vídeo largo sin subtítulos → background job; misma URL para el .txt desde el inicio
    if not has_subtitles and video_duration and video_duration > MAX_DURATION_WHISPER:
        dao = TranscriptionDAO()
        job_id = dao.create_pending(url=url, video_title=video_title, language=lang, username=user, duration_seconds=video_duration)
        base = str(request.base_url).rstrip("/")
        txt_url = f"{base}/export/txt/{job_id}"
        asyncio.create_task(
            _run_transcription_job(job_id, url, lang, include_ts, model)
        )
        return JSONResponse({
            "status": "processing",
            "id": job_id,
            "txt_url": txt_url,
            "message": "Transcripción en segundo plano. Puede tardar hasta 3 horas en vídeos muy largos. Usa el enlace para descargar el .txt cuando esté listo.",
            "video_title": video_title or "Sin título",
        })

    # Flujo síncrono: subtítulos o Whisper (vídeo ≤ 10 min)
    try:
        transcript, err = await asyncio.wait_for(
            asyncio.to_thread(
                subs_service.get_subtitles,
                url,
                language=lang,
                include_timestamps=include_ts,
            ),
            timeout=SUBTITLES_TIMEOUT,
        )
    except asyncio.TimeoutError:
        transcript, err = None, "Tiempo de espera agotado al buscar subtítulos."
    source = "subtitles"

    if not transcript:
        # Vídeo largo sin subtítulos (o no se pudieron obtener) → enviar a segundo plano
        if video_duration and video_duration > MAX_DURATION_WHISPER:
            dao = TranscriptionDAO()
            job_id = dao.create_pending(url=url, video_title=video_title, language=lang, username=user, duration_seconds=video_duration)
            base = str(request.base_url).rstrip("/")
            txt_url = f"{base}/export/txt/{job_id}"
            asyncio.create_task(
                _run_transcription_job(job_id, url, lang, include_ts, model)
            )
            return JSONResponse({
                "status": "processing",
                "id": job_id,
                "txt_url": txt_url,
                "message": "Transcripción en segundo plano. Puede tardar hasta 3 horas en vídeos muy largos. Usa el enlace para descargar el .txt cuando esté listo.",
                "video_title": video_title or "Sin título",
            })

        audio_path, err_audio = await asyncio.to_thread(whisper_service.download_audio, url, COOKIES_FILE)
        if not audio_path:
            return JSONResponse(
                {"error": err_audio or "No se pudo descargar el audio."},
                status_code=400,
            )
        try:
            try:
                transcript = await asyncio.wait_for(
                    asyncio.to_thread(
                        whisper_service.transcribe_audio,
                        audio_path,
                        model_size=model,
                        language=lang if lang != "auto" else None,
                        include_timestamps=include_ts,
                    ),
                    timeout=WHISPER_TIMEOUT,
                )
            except asyncio.TimeoutError:
                return JSONResponse(
                    {"error": "La transcripción con Whisper tardó demasiado. Prueba un vídeo más corto (< 10 min) o uno con subtítulos."},
                    status_code=408,
                )
            source = "whisper"
        finally:
            try:
                if audio_path and os.path.isfile(audio_path):
                    os.remove(audio_path)
            except OSError:
                pass
        if not (transcript or "").strip():
            return JSONResponse(
                {"error": "La transcripción con Whisper no generó texto."},
                status_code=400,
            )

    dao = TranscriptionDAO()
    tid = dao.save(
        url=url,
        video_title=video_title or "",
        source=source,
        transcript=transcript or "",
        language=lang,
        username=user,
        duration_seconds=video_duration or None,
    )
    return JSONResponse({
        "transcript": transcript,
        "video_title": video_title or "Sin título",
        "source": source,
        "id": tid,
    })


@app.post("/upload-cookies")
async def upload_cookies(file: UploadFile = File(...), user=Depends(verify_token)):
    if not file.filename or not file.filename.lower().endswith(".txt"):
        return JSONResponse(
            {"error": "Sube un archivo .txt (cookies en formato Netscape)"},
            status_code=400,
        )
    content = await file.read()
    with open(COOKIES_FILE, "wb") as f:
        f.write(content)
    return JSONResponse({"ok": True, "message": "Cookies guardadas. Ahora podrás descargar videos restringidos."})


def _safe_filename(s: str, max_len: int = 80) -> str:
    s = re.sub(r'[^\w\s\-\.]', "", (s or "").strip()) or "transcripcion"
    return s[:max_len].strip()


@app.get("/export/txt/{id}")
async def export_txt(
    request: Request,
    id: int,
    user=Depends(verify_token),
):
    dao = TranscriptionDAO()
    t = dao.get_by_id(id)
    if not t:
        return JSONResponse({"error": "Transcripción no encontrada."}, status_code=404)
    status = getattr(t, "status", "completed")
    if status == "failed":
        err = getattr(t, "error_message", None) or "La transcripción falló."
        return JSONResponse(
            {"error": err},
            status_code=503,
        )
    if status != "completed" or not (t.transcript or "").strip():
        html = (
            "<!DOCTYPE html><html><head><meta charset='utf-8'><title>En proceso</title></head><body>"
            "<p>Transcripción en curso. Puede tardar hasta 3 horas en vídeos muy largos.</p>"
            "<p>Vuelve más tarde a esta misma dirección para descargar el .txt.</p></body></html>"
        )
        return HTMLResponse(html, status_code=202)
    fn = _safe_filename(t.video_title) + ".txt"
    body = (t.transcript or "").encode("utf-8")
    return StreamingResponse(
        io.BytesIO(body),
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{fn}"'},
    )


@app.get("/export/docx")
async def export_docx(
    request: Request,
    id: int = None,
    user=Depends(verify_token),
):
    if id is None:
        return JSONResponse({"error": "Falta el id de la transcripción."}, status_code=400)
    dao = TranscriptionDAO()
    t = dao.get_by_id(id)
    if not t:
        return JSONResponse({"error": "Transcripción no encontrada."}, status_code=404)
    if getattr(t, "status", "completed") != "completed" or not (t.transcript or "").strip():
        return JSONResponse(
            {"error": "Transcripción aún en proceso."},
            status_code=503,
        )
    doc = Document()
    doc.add_heading(t.video_title or "Transcripción", level=0)
    for block in (t.transcript or "").split("\n"):
        if block.strip():
            doc.add_paragraph(block.strip(), style="Normal")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    fn = _safe_filename(t.video_title) + ".docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fn}"'},
    )


def _build_pdf_reportlab(title: str, transcript: str) -> bytes:
    """Genera PDF con reportlab (UTF-8, fiable)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import cm

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    margin = 2 * cm
    y = height - margin
    line_height = 14
    body_height = 12

    # Título
    title_safe = (title or "Transcripción").replace("\n", " ").strip()[:200]
    if not title_safe:
        title_safe = "Transcripción"
    title_draw = title_safe[:80] + ("..." if len(title_safe) > 80 else "")
    c.setFont("Helvetica-Bold", 14)
    c.drawString(margin, y, title_draw.encode("latin-1", "replace").decode("latin-1"))
    y -= line_height * 1.5

    # Párrafos del transcript
    c.setFont("Helvetica", 10)
    for block in (transcript or "").split("\n"):
        line = block.strip()
        if not line:
            continue
        # Cortar líneas largas para no salirse del margen (aprox 90 chars)
        while line:
            if y < margin + body_height:
                c.showPage()
                c.setFont("Helvetica", 10)
                y = height - margin
            part = line[:90] if len(line) > 90 else line
            c.drawString(margin, y, part.encode("latin-1", "replace").decode("latin-1"))
            y -= body_height
            line = line[90:] if len(line) > 90 else ""

    c.save()
    buf.seek(0)
    return buf.getvalue()


@app.get("/export/pdf")
async def export_pdf(
    request: Request,
    id: int = None,
    user=Depends(verify_token),
):
    if id is None:
        return JSONResponse({"error": "Falta el id de la transcripción."}, status_code=400)
    dao = TranscriptionDAO()
    t = dao.get_by_id(id)
    if not t:
        return JSONResponse({"error": "Transcripción no encontrada."}, status_code=404)
    if getattr(t, "status", "completed") != "completed" or not (t.transcript or "").strip():
        return JSONResponse(
            {"error": "Transcripción aún en proceso."},
            status_code=503,
        )
    title = t.video_title or "Transcripción"
    transcript = (t.transcript or "").strip()
    fn = _safe_filename(title) + ".pdf"

    try:
        pdf_bytes = _build_pdf_reportlab(title, transcript)
    except Exception as e:
        # Fallback a fpdf2 si reportlab falla
        try:
            from fpdf import FPDF
        except ImportError:
            try:
                from fpdf2 import FPDF
            except ImportError:
                return JSONResponse(
                    {"error": "Exportación PDF no disponible (instale reportlab o fpdf2)."},
                    status_code=503,
                )
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=14)
        title_safe = title.encode("latin-1", "replace").decode("latin-1")
        pdf.multi_cell(0, 8, title_safe)
        pdf.ln(4)
        pdf.set_font("Helvetica", size=10)
        for block in transcript.split("\n"):
            if block.strip():
                txt = block.strip().encode("latin-1", "replace").decode("latin-1")
                pdf.multi_cell(0, 6, txt)
        out = pdf.output()
        pdf_bytes = bytes(out) if isinstance(out, (bytearray, memoryview)) else out

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{fn}"'},
    )
